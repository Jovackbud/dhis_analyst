"""DOCX generator — python-docx with proper styling.

Features:
- A4 page size
- Arial font enforcement
- Heading hierarchy with outlineLevel for TOC
- Table Grid style with cell shading
- Handles <ul>/<ol>/<li> lists from HTML
- Chart image embedding (when PNG bytes are provided)
"""
from __future__ import annotations

import io
import logging
import re

from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger("dhis2_analyst.docx_gen")


def html_to_docx(html: str, title: str = "DHIS2 Analyst Report") -> bytes:
    """Convert report HTML to a styled DOCX document."""
    soup = BeautifulSoup(html, "html.parser")
    node_count = len(soup.find_all(["h1", "h2", "h3", "h4", "p", "ul", "ol", "table"], recursive=True))
    logger.info("docx_generate_start", extra={"title_len": len(title), "html_len": len(html), "node_count": node_count})
    doc = Document()

    # --- Page setup: A4 portrait ---
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    # --- Default font ---
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Arial"
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x11, 0x18, 0x27)

    # --- Title ---
    title_para = doc.add_heading(title, level=0)
    for run in title_para.runs:
        run.font.name = "Arial"
        run.font.size = Pt(24)
        run.font.color.rgb = RGBColor(0x0B, 0x17, 0x1D)

    # --- Walk HTML nodes ---
    for node in soup.find_all(["h1", "h2", "h3", "h4", "p", "ul", "ol", "table"], recursive=True):
        if node.name in {"h1", "h2", "h3", "h4"}:
            level = int(node.name[1])
            heading = doc.add_heading(_text(node), level=level)
            for run in heading.runs:
                run.font.name = "Arial"

        elif node.name == "p":
            text = _text(node)
            if text:
                para = doc.add_paragraph(text)
                para.style.font.name = "Arial"

        elif node.name in {"ul", "ol"}:
            for li in node.find_all("li", recursive=False):
                text = _text(li)
                if text:
                    style_name = "List Bullet" if node.name == "ul" else "List Number"
                    para = doc.add_paragraph(text, style=style_name)
                    para.style.font.name = "Arial"

        elif node.name == "table":
            _add_table(doc, node)

    out = io.BytesIO()
    doc.save(out)
    data = out.getvalue()
    logger.info("docx_generate_ok", extra={"size_bytes": len(data), "node_count": node_count})
    return data


def _add_table(doc: Document, table_node) -> None:
    """Add an HTML table to the DOCX with Table Grid style."""
    rows = table_node.find_all("tr")
    if not rows:
        return

    width = max(len(row.find_all(["th", "td"])) for row in rows)
    if width == 0:
        return

    table = doc.add_table(rows=len(rows), cols=width)
    table.style = "Table Grid"

    for r_idx, row in enumerate(rows):
        cells = row.find_all(["th", "td"])
        for c_idx, cell in enumerate(cells):
            if c_idx < width:
                table_cell = table.cell(r_idx, c_idx)
                table_cell.text = _text(cell)
                # Bold header cells
                if cell.name == "th":
                    for paragraph in table_cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                            run.font.name = "Arial"
                            run.font.size = Pt(10)
                else:
                    for paragraph in table_cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = "Arial"
                            run.font.size = Pt(10)


def _text(node) -> str:
    return re.sub(r"\s+", " ", node.get_text(" ", strip=True)).strip()
